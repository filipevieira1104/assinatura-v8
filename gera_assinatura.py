import os
import sys
import time
import shutil
import pyautogui
import cv2
import ctypes
import win32com.client as win32
import customtkinter as ctk
import winreg
import numpy as np
import re
from ctypes import windll
from tkinter import Tk, Entry, Label, Button, messagebox
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import NoSuchElementException
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from PIL import Image

# Função para ajustar caminhos de recursos
def resource_path(relative_path):
    """Retorna o caminho absoluto de um recurso, lidando com executáveis do PyInstaller."""
    if hasattr(sys, '_MEIPASS'):
        # Quando executado como executável
        return os.path.join(sys._MEIPASS, relative_path)
    else:
        # Durante o desenvolvimento
        return os.path.join(os.path.abspath("."), relative_path)

def obter_chromedriver_path():
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, "chromedriver.exe")

def salvar_assinatura(nome, area, email, celular):
    user = os.getlogin()
    signature_dir = os.path.join(os.path.expanduser("~"), "AppData", "Roaming", "Microsoft", "Signatures")
    os.makedirs(signature_dir, exist_ok=True)

    imagem_path = os.path.join(signature_dir, "logo.png")
    original_image_path = resource_path("imagens/logo.png")
    shutil.copy(original_image_path, imagem_path)

    docx_path = os.path.join(signature_dir, "MinhaAssinatura.docx")

    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1)
    table.columns[1].width = Inches(2)

    # Primeira célula (Texto da Assinatura)
    cell1 = table.cell(0, 0)
    paragraph = cell1.paragraphs[0]

    # Nome
    run = paragraph.add_run(f"{nome}\n")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(253, 57, 53)

    # Área
    run = paragraph.add_run(f"{area}\n")
    run.font.size = Pt(12)

    # E-mail
    paragraph.add_run("E-mail: ").font.bold = True
    adicionar_hyperlink(paragraph, f"mailto:{email}", email)
    paragraph.add_run("\n")

    # Celular
    paragraph.add_run("Mobile: ").font.bold = True
    paragraph.add_run(f"{celular}\n").font.size = Pt(12)

    # Site
    paragraph.add_run("Site: ").font.bold = True
    adicionar_hyperlink(paragraph, "https://v8.tech", "https://v8.tech")

    # Segunda célula (Imagem)
    cell2 = table.cell(0, 1)
    cell2.paragraphs[0].add_run().add_picture(imagem_path, width=Inches(2))

    # Adicionar borda à esquerda da célula da imagem
    configurar_borda(cell2, lado="left", tamanho="6", cor="000000")  # 6 é a espessura, "000000" é preto

    # Reduz o espaçamento interno entre as células
    table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(0)  # Espaço depois do texto
    table.cell(0, 1).paragraphs[0].paragraph_format.space_before = Pt(1)  # Espaço antes da imagem

    doc.save(docx_path)

    return docx_path

# Função para adicionar hyperlinks
def adicionar_hyperlink(paragraph, url, texto):
    """
    Adiciona um hyperlink clicável a um parágrafo no Word.
    :param paragraph: Parágrafo onde o link será adicionado.
    :param url: URL ou mailto do hyperlink.
    :param texto: Texto que será exibido no link.
    """
    # Criar o elemento de hyperlink
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Criar o texto do hyperlink
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    new_run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = texto
    new_run.append(text)
    hyperlink.append(new_run)

    # Adicionar o hyperlink ao parágrafo
    paragraph._p.append(hyperlink)

def configurar_borda(celula, lado="left", tamanho="4", cor="000000"):
    """
    Adiciona borda em uma célula específica da tabela.
    :param celula: Objeto célula do python-docx
    :param lado: Lado da borda ('top', 'bottom', 'left', 'right')
    :param tamanho: Espessura da borda em 1/8 de pontos
    :param cor: Cor da borda no formato hexadecimal
    """
    tc = celula._tc
    tcPr = tc.get_or_add_tcPr()
    borda = OxmlElement(f"w:{lado}")
    borda.set(qn("w:val"), "single")
    borda.set(qn("w:sz"), tamanho)  # Espessura da borda
    borda.set(qn("w:space"), "0")   # Espaçamento
    borda.set(qn("w:color"), cor)   # Cor da borda
    tcBorders = OxmlElement("w:tcBorders")
    tcBorders.append(borda)
    tcPr.append(tcBorders)    

def copiar_conteudo_word(docx_path):
    try:
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = True
        doc = word.Documents.Open(docx_path)
        doc.Content.Select()
        doc.Application.Selection.Copy()
        time.sleep(2)
        doc.Close(False)
        word.Quit()
    except Exception as e:
        raise Exception(f"Erro ao copiar conteúdo do Word: {e}")
    time.sleep(2)

# Função para ajustar o template de acordo com o DPI
def ajustar_template(template, escala):
    """
    Ajusta o tamanho do template de acordo com a escala de DPI.

    :param template: Imagem do template (array numpy).
    :param escala: Fator de escala (ex.: 1.25 para 125% de escala).
    :return: Template ajustado ao DPI.
    """
    try:
        # Obter dimensões originais
        altura, largura = template.shape[:2]
        
        # Calcular novas dimensões
        nova_largura = int(largura * escala)
        nova_altura = int(altura * escala)
        
        # Redimensionar template
        template_ajustado = cv2.resize(template, (nova_largura, nova_altura), interpolation=cv2.INTER_AREA)
        return template_ajustado
    except Exception as e:
        raise ValueError(f"Erro ao ajustar template: {e}")

# Função para localizar e clicar em uma imagem na tela
def clicar_na_imagem(relative_template_path):
    try:
        # Obter caminho absoluto do template usando resource_path
        template_path = resource_path(relative_template_path)

        # Verificar se o arquivo existe
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template não encontrado: {template_path}")

        # Capturar a tela
        screenshot = pyautogui.screenshot()
        screenshot = cv2.cvtColor(np.array(screenshot), cv2.COLOR_RGB2BGR)

        # Carregar o template
        template = cv2.imread(template_path, cv2.IMREAD_UNCHANGED)

        if template is None:
            raise ValueError("Erro ao carregar o template. Verifique o caminho e o formato da imagem.")

        # Obter DPI e calcular a escala
        dpi = windll.user32.GetDpiForWindow(windll.user32.GetForegroundWindow())
        escala = dpi / 96.0  # 96 DPI é o padrão

        # Ajustar o template para o DPI
        template = ajustar_template(template, escala)

        # Converter para escala de cinza
        screenshot_gray = cv2.cvtColor(screenshot, cv2.COLOR_BGR2GRAY)
        template_gray = cv2.cvtColor(template, cv2.COLOR_BGR2GRAY)

        # Fazer correspondência de template
        result = cv2.matchTemplate(screenshot_gray, template_gray, cv2.TM_CCOEFF_NORMED)
        _, max_val, _, max_loc = cv2.minMaxLoc(result)

        if max_val > 0.7:
            # Encontrar o centro da correspondência
            click_x, click_y = max_loc
            click_x += template.shape[1] // 2
            click_y += template.shape[0] // 2

            # Realizar o clique
            pyautogui.click(click_x, click_y)
            print(f"Imagem encontrada e clicada em: ({click_x}, {click_y})")
            return True
        else:
            print("Imagem não encontrada.")
            return False
    except Exception as e:
        print(f"Erro ao localizar a imagem: {e}")
        return False

def configurar_assinatura_no_outlook_web(email, senha, docx_path):
    options = webdriver.ChromeOptions()
    options.add_argument("--start-maximized")

    service = Service(obter_chromedriver_path())
    driver = webdriver.Chrome(service=service, options=options)

    try:
        driver.get("https://outlook.office.com/mail/")
        time.sleep(5)

        if "login" in driver.current_url:
            driver.find_element(By.NAME, "loginfmt").send_keys(email + Keys.ENTER)
            time.sleep(2)
            driver.find_element(By.NAME, "passwd").send_keys(senha + Keys.ENTER)
            time.sleep(2)

            # Verificar erro de senha incorreta
            try:
                erro_elemento = driver.find_element(By.ID, "passwordError")
                if erro_elemento.is_displayed():
                    raise Exception("Senha incorreta. Por favor, tente novamente.")
            except NoSuchElementException:
                pass

            messagebox.showinfo("Aguardando Autenticador", "Insira o código do autenticador no navegador e clique em 'OK'.")
            time.sleep(15)
            try:
                driver.find_element(By.ID, "idSIButton9").click()
                time.sleep(5)
            except:
                pass    

        if "mail" not in driver.current_url:
            raise Exception("Erro no login. Verifique as credenciais.")

        driver.find_element(By.ID, "owaSettingsButton").click()
        time.sleep(2)
        WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, "//button[@role='tab' and @value='accounts-category']"))
        ).click()
        time.sleep(3)

        WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, "//button[@role='tab' and @value='signatures-subcategory']"))
        ).click()
        time.sleep(3)

        try:
            excluir_botao = driver.find_element(By.XPATH, "//button[contains(text(),'Excluir')]")
            excluir_botao.click()
            time.sleep(2)
        except:
            print("Nenhuma assinatura existente para excluir.")

        titulo_campo = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, "//input[@placeholder='Editar nome da assinatura']"))
        )
        titulo_campo.clear()
        titulo_campo.send_keys("MinhaAssinatura")

        # Selecionar a assinatura no primeiro menu suspenso
        try:
            assinatura_menu_1 = WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.ID, "fui-r4i"))
            )
            assinatura_menu_1.click()
            time.sleep(1)
            driver.find_element(By.XPATH, "//option[contains(text(), 'MinhaAssinatura')]").click()
            print("Nome da assinatura configurado com sucesso no primeiro seletor.")
        except Exception as e:
            print(f"Erro ao configurar o nome da assinatura no primeiro seletor: {e}")

        # Selecionar a assinatura no segundo menu suspenso
        try:
            assinatura_menu_2 = WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.ID, "fui-r4j"))
            )
            assinatura_menu_2.click()
            time.sleep(1)
            driver.find_element(By.XPATH, "//option[contains(text(), 'MinhaAssinatura')]").click()
            print("Nome da assinatura configurado com sucesso no segundo seletor.")
        except Exception as e:
            print(f"Erro ao configurar o nome da assinatura no segundo seletor: {e}")

        copiar_conteudo_word(docx_path)

        editor_campo = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, "//div[@role='textbox' and @aria-label='Assinatura, pressione Alt+F10 para sair']"))
        )

        editor_campo.click()
        time.sleep(1)
        pyautogui.hotkey('ctrl', 'v')
        time.sleep(2)

        template_path = resource_path("imagens/template_imagem.PNG")  # Caminho ajustado usando resource_path
        if not clicar_na_imagem(template_path):
            print("Não foi possível clicar na área da imagem.")


        inserir_imagem_botao = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, "//button[@aria-label='Inserir imagens embutidas']"))
        )
        inserir_imagem_botao.click()
        time.sleep(2)

        pyautogui.write(os.path.join(os.path.expanduser("~"), "AppData", "Roaming", "Microsoft", "Signatures", "logo.png"))
        pyautogui.press('enter')
        time.sleep(2)

        salvar_botao = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, "//button[@type='button' and contains(text(),'Salvar')]"))
        )
        salvar_botao.click()
        time.sleep(2)
        messagebox.showinfo("Sucesso", f"Assinatura configurada com sucesso para {email}!")

    except Exception as e:
        messagebox.showerror("Erro", f"Erro durante a configuração: {e}")
    finally:
        driver.quit()

def alterar_papel_de_parede():
    try:
        # Caminho do papel de parede
        caminho_papel_de_parede = resource_path("imagens/wallpaper.jpeg")
        if not os.path.exists(caminho_papel_de_parede):
            raise FileNotFoundError(f"Imagem não encontrada: {caminho_papel_de_parede}")

        # Configurar estilo do papel de parede para área de trabalho
        chave_registro = r"Control Panel\Desktop"
        with winreg.OpenKey(winreg.HKEY_CURRENT_USER, chave_registro, 0, winreg.KEY_SET_VALUE) as chave:
            winreg.SetValueEx(chave, "WallpaperStyle", 0, winreg.REG_SZ, "10")  # Preenchido
            winreg.SetValueEx(chave, "TileWallpaper", 0, winreg.REG_SZ, "0")   # Sem mosaico
        ctypes.windll.user32.SystemParametersInfoW(20, 0, caminho_papel_de_parede, 3)

        # Configurar papel de parede da tela de bloqueio
        # chave_lock_screen = r"SOFTWARE\Microsoft\Windows\CurrentVersion\PersonalizationCSP"
        # try:
        #     with winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, chave_lock_screen, 0, winreg.KEY_SET_VALUE) as chave:
        #         winreg.SetValueEx(chave, "LockScreenImagePath", 0, winreg.REG_SZ, caminho_papel_de_parede)
        #         winreg.SetValueEx(chave, "LockScreenImageStatus", 0, winreg.REG_DWORD, 1)
        # except PermissionError:
        #     raise PermissionError("Permissões administrativas necessárias para alterar a tela de bloqueio.")
        
        # Mensagem de sucesso
        messagebox.showinfo("Sucesso", "Papel de parede alterado com sucesso!")

    except FileNotFoundError as e:
        messagebox.showerror("Erro", f"Erro ao alterar o papel de parede: {e}")
    except PermissionError as e:
        messagebox.showerror("Erro", f"Erro ao alterar o papel de parede: {e}")
    except Exception as e:
        messagebox.showerror("Erro", f"Erro inesperado: {e}")

# Função para validar E-mail
def validar_email(email):
    if re.fullmatch(r"^[a-zA-Z0-9_.+-]+@v8\.tech$", email):
        return True
    return False    

# Função para validar número de celular
def validar_celular(celular):
    # Formato: (XX) XXXXX-XXXX ou (XX) XXXX-XXXX
    if re.fullmatch(r"^\(\d{2}\) 9?\d{4}-\d{4}$", celular):
        return True
    return False    
       
def executar_automacao():
    nome = nome_entry.get().strip()
    area = area_var.get()
    email = email_entry.get().strip()
    celular = celular_entry.get().strip()
    senha = senha_entry.get().strip()

    # Verificar se todos os campos estão preenchidos
    if not nome or not area_var or not email or not celular or not senha:
        messagebox.showerror("Erro", "Por favor, preencha todos os campos.")
        return

    # Validação do e-mail
    if not validar_email(email):
        messagebox.showerror("Erro", "E-mail inválido! Favor insira seu e-mail V8.")
        return

    # Validação do celular
    if not validar_celular(celular):
        messagebox.showerror("Erro", "Número de celular inválido! Use o formato (DDD) 91234-5678 ou (DDD) 1234-5678.")
        return

    try:
        # Alterar papel de parede, se necessário
        if alterar_wallpaper_checkbox.get():
            alterar_papel_de_parede()

        # Gerar o caminho para o arquivo de assinatura
        docx_path = salvar_assinatura(nome, area, email, celular)

        # Configurar assinatura no Outlook
        configurar_assinatura_no_outlook_web(email, senha, docx_path)

        # Exibir mensagem de sucesso
        messagebox.showinfo("Sucesso", "Automação concluída com sucesso!")

    except Exception as e:
        # Tratar erros inesperados
        messagebox.showerror("Erro", f"Erro durante a execução da automação: {e}")

# Configuração inicial do CustomTkinter
ctk.set_appearance_mode("System")  # Define o tema do sistema (claro/escuro)
ctk.set_default_color_theme("blue")  # Define o tema padrão

# Função para alternar entre dark e light mode
def toggle_mode():
    current_mode = ctk.get_appearance_mode()
    new_mode = "Light" if current_mode == "Dark" else "Dark"
    ctk.set_appearance_mode(new_mode)
    toggle_button.configure(text=f"Modo: {new_mode}")

def on_select(option):
    print(f'Selecionado: {option}')    

# Criação da janela principal
app = ctk.CTk()
app.title("Configurar Assinatura do Outlook")
app.geometry("470x600")
app.resizable(False, False)  # Desativa redimensionamento da largura e altura

# Adicionando o logo na parte superior
logo_path = resource_path("imagens/v8.png")
logo = ctk.CTkImage(Image.open(logo_path), size=(350, 100))
logo_label = ctk.CTkLabel(app, image=logo, text="")
logo_label.pack(pady=10)

# Campos do formulário
form_frame = ctk.CTkFrame(app)
form_frame.pack(pady=10, padx=10, fill="both", expand=True)
form_frame.grid_columnconfigure(0, weight=1)  # Centraliza os frames no eixo horizontal

# Primeira parte: Dados da assinatura
dados_assinatura_frame = ctk.CTkFrame(form_frame)
dados_assinatura_frame.grid(row=0, column=0, padx=10, pady=10, sticky="n")
ctk.CTkLabel(dados_assinatura_frame, text="Dados da Assinatura", font=("Arial", 16, "bold")).grid(row=0, column=0, columnspan=2, pady=5)

ctk.CTkLabel(dados_assinatura_frame, text="Nome:").grid(row=1, column=0, padx=10, pady=5, sticky="w")
nome_entry = ctk.CTkEntry(dados_assinatura_frame, width=200)
nome_entry.grid(row=1, column=1, padx=10, pady=5)

ctk.CTkLabel(dados_assinatura_frame, text="Celular:").grid(row=2, column=0, padx=10, pady=5, sticky="w")
celular_entry = ctk.CTkEntry(dados_assinatura_frame, width=200)
celular_entry.grid(row=2, column=1, padx=10, pady=5)

ctk.CTkLabel(dados_assinatura_frame, text="Área:").grid(row=3, column=0, padx=10, pady=5, sticky="w")

# Dropdown (CTkOptionMenu) para cargos
area_var = ctk.StringVar(value="Selecione uma opção")  # Variável associada ao dropdown
areas = ["Financeiro", "Comercial", "Delivery", "Marketing", "Tecnologia & Inovação Digital",
         "CO", "Diretor(a)"]  # Lista de opções

setor_dropdown = ctk.CTkOptionMenu(dados_assinatura_frame, variable=area_var, values=areas)
setor_dropdown.grid(row=3, column=1, padx=10, pady=5)

# Segunda parte: Dados do Outlook
dados_outlook_frame = ctk.CTkFrame(form_frame)
dados_outlook_frame.grid(row=1, column=0, padx=10, pady=10, sticky="n")
ctk.CTkLabel(dados_outlook_frame, text="Dados do Outlook", font=("Arial", 16, "bold")).grid(row=0, column=0, columnspan=2, pady=5)

ctk.CTkLabel(dados_outlook_frame, text="E-mail:").grid(row=1, column=0, padx=10, pady=5, sticky="w")
email_entry = ctk.CTkEntry(dados_outlook_frame, width=200)
email_entry.grid(row=1, column=1, padx=10, pady=5)

ctk.CTkLabel(dados_outlook_frame, text="Senha:").grid(row=2, column=0, padx=10, pady=5, sticky="w")
senha_entry = ctk.CTkEntry(dados_outlook_frame, show="*", width=200)
senha_entry.grid(row=2, column=1, padx=10, pady=5)

# Checkbox para alterar papel de parede
alterar_wallpaper_checkbox = ctk.BooleanVar(value=False)
checkbox_widget = ctk.CTkCheckBox(
    form_frame,
    text="Deseja Alterar o papel de parede ?",
    variable=alterar_wallpaper_checkbox
)
checkbox_widget.grid(row=2, column=0, padx=10, pady=(5, 0), sticky="w")

# Botão para executar automação
executar_button = ctk.CTkButton(form_frame, text="Gerar e Configurar Assinatura", command=executar_automacao)
executar_button.grid(row=3, column=0, pady=20)

# Adicionando o botão tipo liga/desliga para alternar dark/light mode
bottom_frame = ctk.CTkFrame(app)
bottom_frame.pack(side="bottom", fill="x", pady=10)

toggle_button = ctk.CTkSwitch(bottom_frame, text="Modo: System", command=toggle_mode, width=150)
toggle_button.pack(side="left", padx=10)

# Loop principal da aplicação
app.mainloop()